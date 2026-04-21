"""Document parsing — unified extraction via Kreuzberg with legacy fallbacks.

Kreuzberg (Rust core) handles 90+ formats including PDF (with OCR),
Office documents, HTML, and images through a single API. Legacy
per-format parsers remain as fallbacks if Kreuzberg is unavailable.
"""
from __future__ import annotations

import csv
import io
import logging
import mimetypes
from dataclasses import dataclass, field
from pathlib import Path
from tempfile import NamedTemporaryFile

logger = logging.getLogger(__name__)

KREUZBERG_SUPPORTED_SUFFIXES = {
    "pdf", "docx", "doc", "pptx", "ppt", "xlsx", "xls", "xlsm",
    "csv", "html", "htm", "md", "txt", "rtf", "odt", "ods", "odp",
    "epub", "xml", "json", "yaml", "yml", "rst", "tex", "log",
    "png", "jpg", "jpeg", "tiff", "bmp", "webp", "gif",
}


@dataclass(slots=True)
class ParsedPage:
    page_no: int
    text_md: str
    elements: dict = field(default_factory=dict)


@dataclass(slots=True)
class ParsedDocument:
    title: str
    mime_type: str
    pages: list[ParsedPage]


def detect_mime_type(filename: str, fallback: str = "application/octet-stream") -> str:
    mime_type, _ = mimetypes.guess_type(filename)
    return mime_type or fallback


def _parse_with_kreuzberg(filename: str, data: bytes, mime_type: str) -> ParsedDocument:
    """Extract text via Kreuzberg — supports 90+ formats, OCR, tables."""
    from kreuzberg import ExtractionConfig, extract_bytes_sync

    suffix = filename.lower().rsplit(".", 1)[-1] if "." in filename else "txt"
    title = filename.rsplit(".", 1)[0] if "." in filename else filename

    config = ExtractionConfig()
    result = extract_bytes_sync(data, mime_type=mime_type, config=config)

    content = result.content or ""
    metadata = result.metadata if hasattr(result, "metadata") else {}
    page_count = 1
    if isinstance(metadata, dict):
        page_count = metadata.get("page_count", 1) or 1

    tables_data = result.tables if hasattr(result, "tables") else []

    if suffix == "pdf" and page_count > 1:
        pages = _split_pdf_pages_heuristic(content, page_count)
    else:
        pages = [ParsedPage(page_no=1, text_md=content)]

    if tables_data:
        for i, table in enumerate(tables_data):
            if hasattr(table, "rows") and table.rows:
                md_table = _table_to_markdown(table.rows)
                if i < len(pages):
                    pages[i].elements["tables"] = pages[i].elements.get("tables", [])
                    pages[i].elements["tables"].append(md_table)

    return ParsedDocument(title=title, mime_type=mime_type, pages=pages)


def _split_pdf_pages_heuristic(content: str, page_count: int) -> list[ParsedPage]:
    """Split extracted text into approximate pages via form-feed or even split."""
    if "\f" in content:
        raw_pages = content.split("\f")
    else:
        lines = content.splitlines(keepends=True)
        chunk_size = max(1, len(lines) // page_count)
        raw_pages = []
        for i in range(0, len(lines), chunk_size):
            raw_pages.append("".join(lines[i : i + chunk_size]))

    pages: list[ParsedPage] = []
    for idx, text in enumerate(raw_pages):
        stripped = text.strip()
        if stripped:
            pages.append(ParsedPage(page_no=idx + 1, text_md=stripped))
    return pages or [ParsedPage(page_no=1, text_md=content.strip())]


def _table_to_markdown(rows: list) -> str:
    """Convert a list of row lists to markdown table format."""
    if not rows:
        return ""
    str_rows = [
        ["" if cell is None else str(cell) for cell in row]
        for row in rows
    ]
    header = "| " + " | ".join(str_rows[0]) + " |"
    separator = "| " + " | ".join("---" for _ in str_rows[0]) + " |"
    body_lines = ["| " + " | ".join(row) + " |" for row in str_rows[1:]]
    return "\n".join([header, separator] + body_lines)


# ── Legacy fallback parsers ──────────────────────────────────────────

def _parse_legacy(filename: str, data: bytes, mime_type: str) -> ParsedDocument:
    """Per-format fallback when Kreuzberg is not available."""
    suffix = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    title = filename.rsplit(".", 1)[0] if "." in filename else filename

    if suffix in {"md", "txt"}:
        text = data.decode("utf-8", errors="replace")
        return ParsedDocument(title=title, mime_type=mime_type, pages=[ParsedPage(page_no=1, text_md=text)])

    if suffix in {"html", "htm"}:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(data.decode("utf-8", errors="replace"), "html.parser")
        text = soup.get_text("\n").strip()
        return ParsedDocument(title=title, mime_type=mime_type, pages=[ParsedPage(page_no=1, text_md=text)])

    if suffix == "pdf":
        try:
            import pdfplumber
            pages: list[ParsedPage] = []
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                for idx, page in enumerate(pdf.pages):
                    text = (page.extract_text() or "").strip()
                    paragraphs: list[dict] = []
                    try:
                        words = page.extract_words(keep_blank_chars=False, use_text_flow=True)
                        if words:
                            current_top = words[0]["top"]
                            current_words: list[dict] = [words[0]]
                            for w in words[1:]:
                                if abs(w["top"] - current_top) > 5.0:
                                    para_text = " ".join(cw["text"] for cw in current_words)
                                    paragraphs.append({
                                        "text": para_text,
                                        "bbox": [
                                            round(min(cw["x0"] for cw in current_words), 2),
                                            round(min(cw["top"] for cw in current_words), 2),
                                            round(max(cw["x1"] for cw in current_words), 2),
                                            round(max(cw["bottom"] for cw in current_words), 2),
                                        ],
                                    })
                                    current_words = [w]
                                    current_top = w["top"]
                                else:
                                    current_words.append(w)
                            if current_words:
                                para_text = " ".join(cw["text"] for cw in current_words)
                                paragraphs.append({
                                    "text": para_text,
                                    "bbox": [
                                        round(min(cw["x0"] for cw in current_words), 2),
                                        round(min(cw["top"] for cw in current_words), 2),
                                        round(max(cw["x1"] for cw in current_words), 2),
                                        round(max(cw["bottom"] for cw in current_words), 2),
                                    ],
                                })
                    except Exception:
                        logger.debug("bbox extraction failed for page %d", idx + 1, exc_info=True)
                    pages.append(ParsedPage(
                        page_no=idx + 1,
                        text_md=text,
                        elements={"paragraphs": paragraphs} if paragraphs else {},
                    ))
            return ParsedDocument(title=title, mime_type=mime_type, pages=pages or [ParsedPage(page_no=1, text_md="")])
        except ImportError:
            logger.warning("pdfplumber not installed, falling back to pypdf")
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        pages_basic = [
            ParsedPage(page_no=i + 1, text_md=(p.extract_text() or "").strip())
            for i, p in enumerate(reader.pages)
        ]
        return ParsedDocument(title=title, mime_type=mime_type, pages=pages_basic or [ParsedPage(page_no=1, text_md="")])

    if suffix == "docx":
        from docx import Document as DocxDocument
        document = DocxDocument(io.BytesIO(data))
        paragraphs = "\n".join(p.text for p in document.paragraphs if p.text.strip())
        return ParsedDocument(title=title, mime_type=mime_type, pages=[ParsedPage(page_no=1, text_md=paragraphs)])

    if suffix in {"xlsx", "xlsm", "xltx", "xltm", "xls"}:
        from openpyxl import load_workbook
        workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        pages: list[ParsedPage] = []
        for index, sheet_name in enumerate(workbook.sheetnames, start=1):
            sheet = workbook[sheet_name]
            rows = [
                ["" if cell is None else str(cell) for cell in row]
                for row in sheet.iter_rows(values_only=True)
            ]
            if not rows:
                rows = [["(empty)"]]
            table = _table_to_markdown(rows[:101])
            pages.append(ParsedPage(page_no=index, text_md=f"## {sheet_name}\n\n{table}", elements={"sheet_name": sheet_name}))
        workbook.close()
        return ParsedDocument(title=title, mime_type=mime_type, pages=pages)

    if suffix == "csv":
        rows = list(csv.reader(io.StringIO(data.decode("utf-8", errors="replace"))))
        if not rows:
            rows = [["(empty)"]]
        table = _table_to_markdown(rows[:101])
        return ParsedDocument(title=title, mime_type=mime_type, pages=[ParsedPage(page_no=1, text_md=table, elements={"sheet_name": "Sheet1"})])

    text = data.decode("utf-8", errors="replace")
    return ParsedDocument(title=title, mime_type=mime_type, pages=[ParsedPage(page_no=1, text_md=text)])


# ── Public entry point ───────────────────────────────────────────────

def parse_document(filename: str, data: bytes, mime_type: str | None = None) -> ParsedDocument:
    effective_mime = mime_type or detect_mime_type(filename)

    try:
        return _parse_with_kreuzberg(filename, data, effective_mime)
    except ImportError:
        logger.info("kreuzberg not installed, using legacy parsers")
    except Exception:
        logger.warning("kreuzberg extraction failed, falling back to legacy", exc_info=True)

    return _parse_legacy(filename, data, effective_mime)
